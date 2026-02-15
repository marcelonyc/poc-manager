"""Enhanced document generation service with charts and comprehensive metrics.

Generates PDF and Markdown reports for POCs that include:
- Executive dashboard with key metrics
- Visual charts (task status, success criteria, progress, timeline, workload, activity)
- Detailed task and task group breakdowns with resources and comments
- Participant information
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Image,
    KeepTogether,
    HRFlowable,
)
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from datetime import datetime, date, timedelta
from typing import Optional, List, Dict, Any
from collections import defaultdict
from sqlalchemy.orm import Session
from app.models.poc import POC, POCParticipant
from app.models.task import POCTask, POCTaskGroup, POCTaskAssignee
from app.models.comment import Comment
from app.models.resource import Resource
from app.models.success_criteria import SuccessCriteria
from app.utils.chart_generator import (
    generate_task_status_pie_chart,
    generate_success_criteria_chart,
    generate_progress_gauge,
    generate_timeline_chart,
    generate_workload_chart,
    generate_activity_chart,
    cleanup_chart_files,
)
from pathlib import Path
import os
import logging

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Enhanced POC document generator with charts and comprehensive metrics."""

    def __init__(self, db: Session, poc: POC):
        self.db = db
        self.poc = poc
        self.tenant = poc.tenant
        self._primary_hex = (
            self.tenant.primary_color
            if self.tenant and self.tenant.primary_color
            else "#0066cc"
        )
        self._secondary_hex = (
            self.tenant.secondary_color
            if self.tenant and self.tenant.secondary_color
            else "#333333"
        )
        self.primary_color = colors.HexColor(self._primary_hex)
        self.secondary_color = colors.HexColor(self._secondary_hex)

        # Caches
        self._all_tasks: Optional[List[POCTask]] = None
        self._all_task_groups: Optional[List[POCTaskGroup]] = None
        self._all_comments: Optional[List[Comment]] = None
        self._metrics: Optional[Dict[str, Any]] = None
        self._chart_files: List[str] = []

    # ------------------------------------------------------------------ #
    #  Data helpers                                                        #
    # ------------------------------------------------------------------ #

    def _get_all_tasks(self) -> List[POCTask]:
        if self._all_tasks is None:
            self._all_tasks = (
                self.db.query(POCTask)
                .filter(POCTask.poc_id == self.poc.id)
                .order_by(POCTask.sort_order)
                .all()
            )
        return self._all_tasks

    def _get_all_task_groups(self) -> List[POCTaskGroup]:
        if self._all_task_groups is None:
            self._all_task_groups = (
                self.db.query(POCTaskGroup)
                .filter(POCTaskGroup.poc_id == self.poc.id)
                .order_by(POCTaskGroup.sort_order)
                .all()
            )
        return self._all_task_groups

    def _get_all_comments(self) -> List[Comment]:
        if self._all_comments is None:
            self._all_comments = (
                self.db.query(Comment)
                .filter(Comment.poc_id == self.poc.id)
                .order_by(Comment.created_at.desc())
                .all()
            )
        return self._all_comments

    def _get_participants(self) -> List[POCParticipant]:
        return (
            self.db.query(POCParticipant)
            .filter(POCParticipant.poc_id == self.poc.id)
            .all()
        )

    def _get_poc_resources(self) -> List[Resource]:
        return (
            self.db.query(Resource)
            .filter(
                Resource.poc_id == self.poc.id,
                Resource.poc_task_id.is_(None),
                Resource.poc_task_group_id.is_(None),
            )
            .all()
        )

    def _get_group_tasks(self, group: POCTaskGroup) -> List[POCTask]:
        """Get POC tasks belonging to a task group."""
        if group.task_group_id:
            from app.models.task import TaskGroup

            template_group = (
                self.db.query(TaskGroup)
                .filter(TaskGroup.id == group.task_group_id)
                .first()
            )
            if template_group and template_group.tasks:
                template_task_ids = [t.id for t in template_group.tasks]
                return (
                    self.db.query(POCTask)
                    .filter(
                        POCTask.poc_id == self.poc.id,
                        POCTask.task_id.in_(template_task_ids),
                    )
                    .order_by(POCTask.sort_order)
                    .all()
                )
        return []

    @staticmethod
    def _comment_author_name(comment: Comment) -> str:
        """Safely get the display name for a comment author."""
        if comment.user:
            return comment.user.full_name or comment.user.email
        if comment.guest_name:
            return f"{comment.guest_name} (guest)"
        return "Unknown"

    # ------------------------------------------------------------------ #
    #  Metrics computation                                                 #
    # ------------------------------------------------------------------ #

    def _compute_metrics(self) -> Dict[str, Any]:
        """Pre-compute all report metrics."""
        if self._metrics is not None:
            return self._metrics

        tasks = self._get_all_tasks()
        groups = self._get_all_task_groups()
        comments = self._get_all_comments()
        participants = self._get_participants()

        # Task status counts
        status_counts: Dict[str, int] = defaultdict(int)
        for t in tasks:
            status_counts[t.status.value] += 1
        total_tasks = len(tasks)
        completed_tasks = status_counts.get("completed", 0)
        completion_pct = (
            round(completed_tasks / total_tasks * 100, 1) if total_tasks else 0
        )

        # Task group status counts
        group_status_counts: Dict[str, int] = defaultdict(int)
        for g in groups:
            val = g.status.value if g.status else "not_started"
            group_status_counts[val] += 1

        # Success criteria
        criteria = self.poc.success_criteria or []
        criteria_met = sum(1 for c in criteria if c.is_met)
        criteria_total = len(criteria)
        criteria_pct = (
            round(criteria_met / criteria_total * 100, 1)
            if criteria_total
            else 0
        )

        criteria_data = [
            {
                "title": c.title,
                "target_value": c.target_value,
                "achieved_value": c.achieved_value,
                "is_met": c.is_met,
                "importance_level": c.importance_level or 3,
            }
            for c in criteria
        ]

        # Timeline
        days_total = None
        days_elapsed = None
        days_remaining = None
        if self.poc.start_date and self.poc.end_date:
            today = date.today()
            start = (
                self.poc.start_date.date()
                if isinstance(self.poc.start_date, datetime)
                else self.poc.start_date
            )
            end = (
                self.poc.end_date.date()
                if isinstance(self.poc.end_date, datetime)
                else self.poc.end_date
            )
            days_total = (end - start).days or 1
            days_elapsed = max(0, (today - start).days)
            days_remaining = max(0, (end - today).days)

        # Workload per assignee
        assignee_workload: Dict[str, Dict[str, int]] = {}
        for task in tasks:
            assignees = (
                self.db.query(POCTaskAssignee)
                .filter(POCTaskAssignee.poc_task_id == task.id)
                .all()
            )
            for a in assignees:
                if a.participant and a.participant.user:
                    name = (
                        a.participant.user.full_name
                        or a.participant.user.email
                    )
                    if name not in assignee_workload:
                        assignee_workload[name] = defaultdict(int)
                    assignee_workload[name][task.status.value] += 1

        # Activity over time (comments grouped by week)
        activity_data: List[Dict[str, Any]] = []
        if comments:
            week_counts: Dict[str, int] = defaultdict(int)
            for c in comments:
                week_start = c.created_at - timedelta(
                    days=c.created_at.weekday()
                )
                key = week_start.strftime("%b %d")
                week_counts[key] += 1
            # Sort by actual date and take last 12 weeks
            sorted_weeks = sorted(
                week_counts.items(),
                key=lambda x: datetime.strptime(
                    x[0] + f" {datetime.now().year}", "%b %d %Y"
                ),
            )
            activity_data = [
                {"period": k, "count": v} for k, v in sorted_weeks[-12:]
            ]

        # Timeline chart data
        timeline_items = []
        for g in groups:
            start_d = None
            end_d = None
            if hasattr(g, "start_date") and g.start_date:
                start_d = (
                    g.start_date.date()
                    if isinstance(g.start_date, datetime)
                    else g.start_date
                )
            if hasattr(g, "due_date") and g.due_date:
                end_d = (
                    g.due_date.date()
                    if isinstance(g.due_date, datetime)
                    else g.due_date
                )
            if start_d or end_d:
                timeline_items.append(
                    {
                        "title": g.title,
                        "start_date": start_d,
                        "due_date": end_d,
                        "status": (
                            g.status.value if g.status else "not_started"
                        ),
                        "is_group": True,
                    }
                )
        for t in tasks:
            start_d = None
            end_d = None
            if t.start_date:
                start_d = (
                    t.start_date.date()
                    if isinstance(t.start_date, datetime)
                    else t.start_date
                )
            if t.due_date:
                end_d = (
                    t.due_date.date()
                    if isinstance(t.due_date, datetime)
                    else t.due_date
                )
            if start_d or end_d:
                timeline_items.append(
                    {
                        "title": t.title,
                        "start_date": start_d,
                        "due_date": end_d,
                        "status": t.status.value,
                        "is_group": False,
                    }
                )

        self._metrics = {
            "total_tasks": total_tasks,
            "completed_tasks": completed_tasks,
            "completion_pct": completion_pct,
            "status_counts": dict(status_counts),
            "group_status_counts": dict(group_status_counts),
            "total_groups": len(groups),
            "criteria_met": criteria_met,
            "criteria_total": criteria_total,
            "criteria_pct": criteria_pct,
            "criteria_data": criteria_data,
            "total_comments": len(comments),
            "internal_comments": sum(1 for c in comments if c.is_internal),
            "total_participants": len(participants),
            "days_total": days_total,
            "days_elapsed": days_elapsed,
            "days_remaining": days_remaining,
            "assignee_workload": assignee_workload,
            "activity_data": activity_data,
            "timeline_items": timeline_items,
        }
        return self._metrics

    # ------------------------------------------------------------------ #
    #  Chart generation                                                    #
    # ------------------------------------------------------------------ #

    def _generate_charts(
        self, as_base64: bool = False
    ) -> Dict[str, Optional[str]]:
        """Generate all charts and return paths/data URIs."""
        m = self._compute_metrics()
        charts: Dict[str, Optional[str]] = {}

        try:
            charts["progress_gauge"] = generate_progress_gauge(
                m["completion_pct"],
                self.poc.overall_success_score,
                self._primary_hex,
                as_base64,
            )
        except Exception as e:
            logger.warning("Failed to generate progress gauge: %s", e)
            charts["progress_gauge"] = None

        try:
            charts["task_status"] = generate_task_status_pie_chart(
                m["status_counts"],
                self._primary_hex,
                as_base64,
            )
        except Exception as e:
            logger.warning("Failed to generate task status chart: %s", e)
            charts["task_status"] = None

        try:
            charts["success_criteria"] = generate_success_criteria_chart(
                m["criteria_data"],
                self._primary_hex,
                as_base64,
            )
        except Exception as e:
            logger.warning("Failed to generate success criteria chart: %s", e)
            charts["success_criteria"] = None

        poc_start = None
        poc_end = None
        if self.poc.start_date:
            poc_start = (
                self.poc.start_date.date()
                if isinstance(self.poc.start_date, datetime)
                else self.poc.start_date
            )
        if self.poc.end_date:
            poc_end = (
                self.poc.end_date.date()
                if isinstance(self.poc.end_date, datetime)
                else self.poc.end_date
            )

        try:
            charts["timeline"] = generate_timeline_chart(
                m["timeline_items"],
                poc_start,
                poc_end,
                self._primary_hex,
                as_base64,
            )
        except Exception as e:
            logger.warning("Failed to generate timeline chart: %s", e)
            charts["timeline"] = None

        try:
            charts["workload"] = generate_workload_chart(
                m["assignee_workload"],
                self._primary_hex,
                as_base64,
            )
        except Exception as e:
            logger.warning("Failed to generate workload chart: %s", e)
            charts["workload"] = None

        try:
            charts["activity"] = generate_activity_chart(
                m["activity_data"],
                self._primary_hex,
                as_base64,
            )
        except Exception as e:
            logger.warning("Failed to generate activity chart: %s", e)
            charts["activity"] = None

        # Track file paths for cleanup
        for v in charts.values():
            if v and not v.startswith("data:"):
                self._chart_files.append(v)

        return charts

    # ------------------------------------------------------------------ #
    #  PDF Generation                                                      #
    # ------------------------------------------------------------------ #

    def generate_pdf(self, output_path: str):
        """Generate a comprehensive PDF report with charts."""
        m = self._compute_metrics()
        charts = self._generate_charts(as_base64=False)

        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                leftMargin=0.75 * inch,
                rightMargin=0.75 * inch,
                topMargin=0.6 * inch,
                bottomMargin=0.6 * inch,
            )
            styles = getSampleStyleSheet()
            story = []

            # --- Custom styles ---
            title_style = ParagraphStyle(
                "CoverTitle",
                parent=styles["Title"],
                fontSize=28,
                textColor=self.primary_color,
                alignment=TA_CENTER,
                spaceAfter=10,
            )
            subtitle_style = ParagraphStyle(
                "CoverSubtitle",
                parent=styles["Heading2"],
                fontSize=14,
                textColor=colors.HexColor("#6B7280"),
                alignment=TA_CENTER,
                spaceAfter=6,
            )
            section_style = ParagraphStyle(
                "SectionHeading",
                parent=styles["Heading2"],
                fontSize=16,
                textColor=self.primary_color,
                spaceBefore=18,
                spaceAfter=10,
            )
            subsection_style = ParagraphStyle(
                "SubSection",
                parent=styles["Heading3"],
                fontSize=12,
                textColor=self.secondary_color,
                spaceBefore=10,
                spaceAfter=6,
            )
            body_style = styles["BodyText"]
            body_indent = ParagraphStyle(
                "BodyIndent",
                parent=body_style,
                leftIndent=20,
            )

            # =============================================
            #  COVER PAGE
            # =============================================
            story.append(Spacer(1, 1.2 * inch))

            # Customer logo
            if self.poc.customer_logo_url:
                logo_path = self._get_logo_path(self.poc.customer_logo_url)
                if logo_path and os.path.exists(logo_path):
                    try:
                        logo = Image(
                            logo_path,
                            width=2 * inch,
                            height=2 * inch,
                            kind="proportional",
                        )
                        logo.hAlign = "CENTER"
                        story.append(logo)
                        story.append(Spacer(1, 0.3 * inch))
                    except Exception:
                        pass

            # Customer company name
            story.append(
                Paragraph(self.poc.customer_company_name, subtitle_style)
            )
            story.append(Spacer(1, 0.2 * inch))

            # POC Title
            story.append(Paragraph(self.poc.title, title_style))
            story.append(Spacer(1, 0.15 * inch))

            # Status badge
            status_text = self.poc.status.value.replace("_", " ").title()
            story.append(
                Paragraph(
                    f'<font color="#6B7280">Status:</font> <b>{status_text}</b>',
                    ParagraphStyle(
                        "StatusLine",
                        parent=body_style,
                        alignment=TA_CENTER,
                        fontSize=12,
                    ),
                )
            )

            # Date range
            start_str = (
                str(self.poc.start_date) if self.poc.start_date else "TBD"
            )
            end_str = str(self.poc.end_date) if self.poc.end_date else "TBD"
            story.append(
                Paragraph(
                    f'<font color="#6B7280">{start_str}  &mdash;  {end_str}</font>',
                    ParagraphStyle(
                        "DateRange",
                        parent=body_style,
                        alignment=TA_CENTER,
                        fontSize=11,
                    ),
                )
            )
            story.append(Spacer(1, 0.3 * inch))

            # Generation timestamp
            story.append(
                Paragraph(
                    f'<font color="#9CA3AF" size="9">Generated on '
                    f'{datetime.now().strftime("%B %d, %Y at %H:%M")}</font>',
                    ParagraphStyle(
                        "GenDate", parent=body_style, alignment=TA_CENTER
                    ),
                )
            )
            story.append(PageBreak())

            # =============================================
            #  EXECUTIVE DASHBOARD
            # =============================================
            story.append(Paragraph("Executive Dashboard", section_style))
            story.append(
                HRFlowable(
                    width="100%",
                    thickness=1,
                    color=self.primary_color,
                    spaceAfter=10,
                )
            )

            # Key metrics cards as a table
            score_text = (
                f"{self.poc.overall_success_score}/100"
                if self.poc.overall_success_score
                else "N/A"
            )
            days_text = (
                f"{m['days_remaining']}"
                if m["days_remaining"] is not None
                else "N/A"
            )

            metric_data = [
                [
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f'<b>{m["completion_pct"]}%</b></font><br/>'
                        f'<font size="8" color="#6B7280">Task Completion</font>',
                        ParagraphStyle("M1", alignment=TA_CENTER),
                    ),
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f'<b>{m["completed_tasks"]}/{m["total_tasks"]}</b></font><br/>'
                        f'<font size="8" color="#6B7280">Tasks Done</font>',
                        ParagraphStyle("M2", alignment=TA_CENTER),
                    ),
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f'<b>{m["criteria_met"]}/{m["criteria_total"]}</b></font><br/>'
                        f'<font size="8" color="#6B7280">Criteria Met</font>',
                        ParagraphStyle("M3", alignment=TA_CENTER),
                    ),
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f"<b>{score_text}</b></font><br/>"
                        f'<font size="8" color="#6B7280">Success Score</font>',
                        ParagraphStyle("M4", alignment=TA_CENTER),
                    ),
                ],
                [
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f"<b>{days_text}</b></font><br/>"
                        f'<font size="8" color="#6B7280">Days Remaining</font>',
                        ParagraphStyle("M5", alignment=TA_CENTER),
                    ),
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f'<b>{m["total_groups"]}</b></font><br/>'
                        f'<font size="8" color="#6B7280">Task Groups</font>',
                        ParagraphStyle("M6", alignment=TA_CENTER),
                    ),
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f'<b>{m["total_comments"]}</b></font><br/>'
                        f'<font size="8" color="#6B7280">Comments</font>',
                        ParagraphStyle("M7", alignment=TA_CENTER),
                    ),
                    Paragraph(
                        f'<font size="20" color="{self._primary_hex}">'
                        f'<b>{m["total_participants"]}</b></font><br/>'
                        f'<font size="8" color="#6B7280">Participants</font>',
                        ParagraphStyle("M8", alignment=TA_CENTER),
                    ),
                ],
            ]

            col_w = 1.7 * inch
            metric_table = Table(metric_data, colWidths=[col_w] * 4)
            metric_table.setStyle(
                TableStyle(
                    [
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        (
                            "BOX",
                            (0, 0),
                            (-1, -1),
                            0.5,
                            colors.HexColor("#E5E7EB"),
                        ),
                        (
                            "INNERGRID",
                            (0, 0),
                            (-1, -1),
                            0.5,
                            colors.HexColor("#E5E7EB"),
                        ),
                        ("TOPPADDING", (0, 0), (-1, -1), 12),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                        (
                            "BACKGROUND",
                            (0, 0),
                            (-1, -1),
                            colors.HexColor("#F9FAFB"),
                        ),
                    ]
                )
            )
            story.append(metric_table)
            story.append(Spacer(1, 0.3 * inch))

            # --- Charts ---
            if charts.get("progress_gauge"):
                try:
                    img = Image(
                        charts["progress_gauge"],
                        width=5.5 * inch,
                        height=2.5 * inch,
                    )
                    img.hAlign = "CENTER"
                    story.append(img)
                    story.append(Spacer(1, 0.2 * inch))
                except Exception:
                    pass

            if charts.get("task_status"):
                try:
                    img = Image(
                        charts["task_status"],
                        width=4.5 * inch,
                        height=3.2 * inch,
                    )
                    img.hAlign = "CENTER"
                    story.append(img)
                    story.append(Spacer(1, 0.2 * inch))
                except Exception:
                    pass

            if charts.get("success_criteria"):
                try:
                    n_criteria = len(m["criteria_data"])
                    chart_h = max(2.5, n_criteria * 0.6 + 1)
                    img = Image(
                        charts["success_criteria"],
                        width=6 * inch,
                        height=chart_h * inch,
                    )
                    img.hAlign = "CENTER"
                    story.append(img)
                    story.append(Spacer(1, 0.2 * inch))
                except Exception:
                    pass

            story.append(PageBreak())

            # Timeline chart
            if charts.get("timeline"):
                story.append(Paragraph("POC Timeline", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )
                try:
                    n_items = len(m["timeline_items"])
                    chart_h = max(2.5, n_items * 0.4 + 1.5)
                    img = Image(
                        charts["timeline"],
                        width=6.5 * inch,
                        height=chart_h * inch,
                    )
                    img.hAlign = "CENTER"
                    story.append(img)
                    story.append(Spacer(1, 0.3 * inch))
                except Exception:
                    pass

            # Workload chart
            if charts.get("workload"):
                story.append(Paragraph("Team Workload", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )
                try:
                    n_assignees = len(m["assignee_workload"])
                    chart_h = max(2.5, n_assignees * 0.5 + 1)
                    img = Image(
                        charts["workload"],
                        width=6 * inch,
                        height=chart_h * inch,
                    )
                    img.hAlign = "CENTER"
                    story.append(img)
                    story.append(Spacer(1, 0.3 * inch))
                except Exception:
                    pass

            # Activity chart
            if charts.get("activity"):
                story.append(Paragraph("Activity Over Time", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )
                try:
                    img = Image(
                        charts["activity"],
                        width=6 * inch,
                        height=2.8 * inch,
                    )
                    img.hAlign = "CENTER"
                    story.append(img)
                    story.append(Spacer(1, 0.3 * inch))
                except Exception:
                    pass

            story.append(PageBreak())

            # =============================================
            #  DESCRIPTION / EXECUTIVE SUMMARY / OBJECTIVES
            # =============================================
            if (
                self.poc.description
                or self.poc.executive_summary
                or self.poc.objectives
            ):
                story.append(Paragraph("POC Overview", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )

                if self.poc.executive_summary:
                    story.append(
                        Paragraph("Executive Summary", subsection_style)
                    )
                    story.append(
                        Paragraph(self.poc.executive_summary, body_style)
                    )
                    story.append(Spacer(1, 0.15 * inch))

                if self.poc.description:
                    story.append(Paragraph("Description", subsection_style))
                    story.append(Paragraph(self.poc.description, body_style))
                    story.append(Spacer(1, 0.15 * inch))

                if self.poc.objectives:
                    story.append(Paragraph("Objectives", subsection_style))
                    story.append(Paragraph(self.poc.objectives, body_style))
                    story.append(Spacer(1, 0.15 * inch))

            # Products
            if self.poc.products:
                story.append(Paragraph("Products", subsection_style))
                for product in self.poc.products:
                    story.append(
                        Paragraph(f"&bull; {product.name}", body_style)
                    )
                story.append(Spacer(1, 0.15 * inch))

            # =============================================
            #  SUCCESS CRITERIA TABLE
            # =============================================
            criteria = self.poc.success_criteria
            if criteria:
                story.append(Paragraph("Success Criteria", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )

                criteria_header = [
                    [
                        "#",
                        "Criteria",
                        "Target",
                        "Achieved",
                        "Importance",
                        "Met",
                    ]
                ]
                criteria_rows = []
                for i, c in enumerate(criteria, 1):
                    stars = "*" * (c.importance_level or 3)
                    met_icon = "Yes" if c.is_met else "No"
                    criteria_rows.append(
                        [
                            str(i),
                            c.title,
                            c.target_value or "N/A",
                            c.achieved_value or "N/A",
                            stars,
                            met_icon,
                        ]
                    )

                criteria_table = Table(
                    criteria_header + criteria_rows,
                    colWidths=[
                        0.4 * inch,
                        2.2 * inch,
                        1.2 * inch,
                        1.2 * inch,
                        0.9 * inch,
                        0.5 * inch,
                    ],
                )
                criteria_table.setStyle(
                    TableStyle(
                        [
                            (
                                "BACKGROUND",
                                (0, 0),
                                (-1, 0),
                                self.primary_color,
                            ),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            (
                                "FONTNAME",
                                (0, 0),
                                (-1, 0),
                                "Helvetica-Bold",
                            ),
                            ("FONTSIZE", (0, 0), (-1, -1), 9),
                            ("ALIGN", (0, 0), (0, -1), "CENTER"),
                            ("ALIGN", (4, 0), (5, -1), "CENTER"),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            (
                                "GRID",
                                (0, 0),
                                (-1, -1),
                                0.5,
                                colors.HexColor("#E5E7EB"),
                            ),
                            (
                                "ROWBACKGROUNDS",
                                (0, 1),
                                (-1, -1),
                                [colors.white, colors.HexColor("#F9FAFB")],
                            ),
                        ]
                    )
                )
                story.append(criteria_table)
                story.append(Spacer(1, 0.3 * inch))

            # =============================================
            #  INDIVIDUAL TASKS
            # =============================================
            tasks = self._get_all_tasks()
            if tasks:
                story.append(Paragraph("Tasks", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )

                for task in tasks:
                    story.extend(
                        self._pdf_render_task(
                            task,
                            styles,
                            body_style,
                            body_indent,
                            subsection_style,
                        )
                    )

            # =============================================
            #  TASK GROUPS
            # =============================================
            groups = self._get_all_task_groups()
            if groups:
                story.append(PageBreak())
                story.append(Paragraph("Task Groups", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )

                for group in groups:
                    status_val = (
                        group.status.value if group.status else "not_started"
                    )
                    status_label = status_val.replace("_", " ").title()
                    group_title = f"{group.title}  &mdash;  {status_label}"

                    group_heading_style = ParagraphStyle(
                        "GrpTitle",
                        parent=styles["Heading3"],
                        textColor=self.primary_color,
                        fontSize=13,
                        spaceBefore=14,
                    )
                    story.append(Paragraph(group_title, group_heading_style))

                    if group.description:
                        story.append(Paragraph(group.description, body_style))
                        story.append(Spacer(1, 0.1 * inch))

                    # Group resources
                    group_resources = (
                        self.db.query(Resource)
                        .filter(Resource.poc_task_group_id == group.id)
                        .all()
                    )
                    if group_resources:
                        story.append(
                            Paragraph("<b>Group Resources:</b>", body_style)
                        )
                        for res in group_resources:
                            res_text = (
                                f"&bull; <b>{res.title}</b>"
                                f" ({res.resource_type.value})"
                            )
                            if res.description:
                                res_text += f" &mdash; {res.description}"
                            if (
                                res.resource_type.value == "LINK"
                                and res.content
                            ):
                                res_text += (
                                    f'<br/>  <a href="{res.content}"'
                                    f' color="blue">{res.content}</a>'
                                )
                            story.append(Paragraph(res_text, body_indent))
                        story.append(Spacer(1, 0.1 * inch))

                    # Group comments
                    group_comments = (
                        self.db.query(Comment)
                        .filter(Comment.poc_task_group_id == group.id)
                        .order_by(Comment.created_at.desc())
                        .limit(5)
                        .all()
                    )
                    if group_comments:
                        story.append(
                            Paragraph("<b>Group Comments:</b>", body_style)
                        )
                        for c in reversed(group_comments):
                            author = self._comment_author_name(c)
                            dt = c.created_at.strftime("%Y-%m-%d %H:%M")
                            vis = "Internal" if c.is_internal else "Public"
                            story.append(
                                Paragraph(
                                    f"&bull; <b>{author}</b> ({vis})"
                                    f" &mdash; {dt}<br/>  {c.content}",
                                    body_indent,
                                )
                            )
                        story.append(Spacer(1, 0.1 * inch))

                    # Tasks in group
                    group_tasks = self._get_group_tasks(group)
                    if group_tasks:
                        for task in group_tasks:
                            story.extend(
                                self._pdf_render_task(
                                    task,
                                    styles,
                                    body_indent,
                                    body_indent,
                                    subsection_style,
                                    indent=True,
                                )
                            )

                    story.append(Spacer(1, 0.15 * inch))

            # =============================================
            #  POC-LEVEL RESOURCES
            # =============================================
            poc_resources = self._get_poc_resources()
            if poc_resources:
                story.append(Paragraph("POC Resources", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )
                for resource in poc_resources:
                    story.append(
                        Paragraph(
                            f"<b>{resource.title}</b>"
                            f" ({resource.resource_type.value})",
                            subsection_style,
                        )
                    )
                    if resource.description:
                        story.append(
                            Paragraph(resource.description, body_style)
                        )
                    if (
                        resource.resource_type.value == "LINK"
                        and resource.content
                    ):
                        story.append(
                            Paragraph(
                                f'<a href="{resource.content}"'
                                f' color="blue">{resource.content}</a>',
                                body_style,
                            )
                        )
                    elif resource.content:
                        story.append(Paragraph(resource.content, body_style))
                    story.append(Spacer(1, 0.1 * inch))

            # =============================================
            #  PARTICIPANTS
            # =============================================
            participants = self._get_participants()
            if participants:
                story.append(Paragraph("Participants", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )

                p_header = [["Name", "Email", "Role", "Joined"]]
                p_rows = []
                for p in participants:
                    if p.user:
                        name = p.user.full_name or p.user.email
                        email = p.user.email
                        role_parts = []
                        if p.is_sales_engineer:
                            role_parts.append("Sales Engineer")
                        if p.is_customer:
                            role_parts.append("Customer")
                        role = (
                            ", ".join(role_parts)
                            if role_parts
                            else "Participant"
                        )
                        joined = (
                            p.joined_at.strftime("%Y-%m-%d")
                            if p.joined_at
                            else "N/A"
                        )
                        p_rows.append([name, email, role, joined])

                if p_rows:
                    p_table = Table(
                        p_header + p_rows,
                        colWidths=[
                            1.8 * inch,
                            2.2 * inch,
                            1.3 * inch,
                            1 * inch,
                        ],
                    )
                    p_table.setStyle(
                        TableStyle(
                            [
                                (
                                    "BACKGROUND",
                                    (0, 0),
                                    (-1, 0),
                                    self.primary_color,
                                ),
                                (
                                    "TEXTCOLOR",
                                    (0, 0),
                                    (-1, 0),
                                    colors.white,
                                ),
                                (
                                    "FONTNAME",
                                    (0, 0),
                                    (-1, 0),
                                    "Helvetica-Bold",
                                ),
                                ("FONTSIZE", (0, 0), (-1, -1), 9),
                                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                                (
                                    "GRID",
                                    (0, 0),
                                    (-1, -1),
                                    0.5,
                                    colors.HexColor("#E5E7EB"),
                                ),
                                (
                                    "ROWBACKGROUNDS",
                                    (0, 1),
                                    (-1, -1),
                                    [
                                        colors.white,
                                        colors.HexColor("#F9FAFB"),
                                    ],
                                ),
                            ]
                        )
                    )
                    story.append(p_table)

            # =============================================
            #  POC-LEVEL COMMENTS
            # =============================================
            poc_comments = (
                self.db.query(Comment)
                .filter(
                    Comment.poc_id == self.poc.id,
                    Comment.poc_task_id.is_(None),
                    Comment.poc_task_group_id.is_(None),
                )
                .order_by(Comment.created_at.desc())
                .limit(10)
                .all()
            )
            if poc_comments:
                story.append(Paragraph("Recent POC Comments", section_style))
                story.append(
                    HRFlowable(
                        width="100%",
                        thickness=1,
                        color=self.primary_color,
                        spaceAfter=10,
                    )
                )
                for c in reversed(poc_comments):
                    author = self._comment_author_name(c)
                    dt = c.created_at.strftime("%Y-%m-%d %H:%M")
                    vis = "Internal" if c.is_internal else "Public"
                    story.append(
                        Paragraph(
                            f"&bull; <b>{author}</b> ({vis})"
                            f" &mdash; {dt}<br/>  {c.content}",
                            body_style,
                        )
                    )
                story.append(Spacer(1, 0.1 * inch))

            # Build PDF
            doc.build(story)

        finally:
            cleanup_chart_files(self._chart_files)

        return output_path

    def _pdf_render_task(
        self,
        task,
        styles,
        body_style,
        body_indent,
        subsection_style,
        indent=False,
    ):
        """Render a single task to a list of PDF flowable elements."""
        elements = []
        status_emoji = {
            "not_started": "[ ]",
            "in_progress": "[~]",
            "completed": "[x]",
            "blocked": "[!]",
        }
        marker = status_emoji.get(task.status.value, "[ ]")
        status_label = task.status.value.replace("_", " ").title()
        task_title = f"{marker} {task.title}  &mdash;  {status_label}"

        left = 20 if indent else 0
        task_style = ParagraphStyle(
            "TaskTitleStyle",
            parent=subsection_style,
            leftIndent=left,
        )
        elements.append(Paragraph(task_title, task_style))

        tb = ParagraphStyle("TaskBody", parent=body_style, leftIndent=left)

        if task.description:
            elements.append(Paragraph(task.description, tb))

        # Dates
        date_parts = []
        if task.start_date:
            date_parts.append(f"Start: {task.start_date}")
        if task.due_date:
            date_parts.append(f"Due: {task.due_date}")
        if task.completed_at:
            date_parts.append(f"Completed: {task.completed_at}")
        if date_parts:
            elements.append(
                Paragraph(
                    f'<font color="#6B7280">'
                    f'{" | ".join(date_parts)}</font>',
                    tb,
                )
            )

        # Assignees
        assignees = (
            self.db.query(POCTaskAssignee)
            .filter(POCTaskAssignee.poc_task_id == task.id)
            .all()
        )
        if assignees:
            names = []
            for a in assignees:
                if a.participant and a.participant.user:
                    name = (
                        a.participant.user.full_name
                        or a.participant.user.email
                    )
                    names.append(name)
            if names:
                elements.append(
                    Paragraph(f"<b>Assigned to:</b> {', '.join(names)}", tb)
                )

        # Resources
        task_resources = (
            self.db.query(Resource)
            .filter(Resource.poc_task_id == task.id)
            .all()
        )
        if task_resources:
            elements.append(Paragraph("<b>Resources:</b>", tb))
            for res in task_resources:
                res_text = (
                    f"&bull; <b>{res.title}</b>"
                    f" ({res.resource_type.value})"
                )
                if res.description:
                    res_text += f" &mdash; {res.description}"
                if res.resource_type.value == "LINK" and res.content:
                    res_text += (
                        f'<br/>  <a href="{res.content}"'
                        f' color="blue">{res.content}</a>'
                    )
                elements.append(Paragraph(res_text, tb))

        # Success criteria linked to this task
        if hasattr(task, "task_criteria") and task.task_criteria:
            elements.append(Paragraph("<b>Success Criteria:</b>", tb))
            for tc in task.task_criteria:
                sc = (
                    tc.success_criteria
                    if hasattr(tc, "success_criteria")
                    else None
                )
                if sc:
                    met = "Yes" if sc.is_met else "No"
                    elements.append(
                        Paragraph(f"&bull; [{met}] {sc.title}", tb)
                    )

        # Comments
        comments = (
            self.db.query(Comment)
            .filter(Comment.poc_task_id == task.id)
            .order_by(Comment.created_at.desc())
            .limit(5)
            .all()
        )
        if comments:
            elements.append(Paragraph("<b>Latest Comments:</b>", tb))
            for c in reversed(comments):
                author = self._comment_author_name(c)
                dt = c.created_at.strftime("%Y-%m-%d %H:%M")
                vis = "Internal" if c.is_internal else "Public"
                elements.append(
                    Paragraph(
                        f"&bull; <b>{author}</b> ({vis})"
                        f" &mdash; {dt}<br/>  {c.content}",
                        tb,
                    )
                )

        elements.append(Spacer(1, 0.15 * inch))
        return elements

    # ------------------------------------------------------------------ #
    #  Markdown Generation                                                 #
    # ------------------------------------------------------------------ #

    def generate_markdown(self, output_path: str):
        """Generate a comprehensive Markdown report with embedded charts."""
        m = self._compute_metrics()
        charts = self._generate_charts(as_base64=True)

        md: List[str] = []

        # =============================================
        #  HEADER / COVER
        # =============================================
        if self.poc.customer_logo_url:
            md.append('<div align="center">')
            md.append(
                f'<img src="{self.poc.customer_logo_url}"'
                f' alt="Customer Logo" width="200"/>'
            )
            md.append("</div>\n")

        md.append('<div align="center">')
        md.append(
            f'<h2 style="color: {self._secondary_hex};">'
            f"{self.poc.customer_company_name}</h2>"
        )
        md.append("</div>\n")

        md.append(
            f'<h1 style="color: {self._primary_hex};">'
            f"{self.poc.title}</h1>\n"
        )

        status_text = self.poc.status.value.replace("_", " ").title()
        start_str = str(self.poc.start_date) if self.poc.start_date else "TBD"
        end_str = str(self.poc.end_date) if self.poc.end_date else "TBD"
        md.append(
            f"**Status:** {status_text} &nbsp;|&nbsp;"
            f" **Period:** {start_str} — {end_str}\n"
        )
        md.append(
            f"<sub>Generated on"
            f' {datetime.now().strftime("%B %d, %Y at %H:%M")}</sub>\n'
        )
        md.append("---\n")

        # =============================================
        #  EXECUTIVE DASHBOARD
        # =============================================
        md.append(
            f'<h2 style="color: {self._primary_hex};">'
            f"Executive Dashboard</h2>\n"
        )

        score_text = (
            f"{self.poc.overall_success_score}/100"
            if self.poc.overall_success_score
            else "N/A"
        )
        days_text = (
            str(m["days_remaining"])
            if m["days_remaining"] is not None
            else "N/A"
        )

        md.append("| Metric | Value |")
        md.append("|--------|-------|")
        md.append(
            f"| **Task Completion** |"
            f" {m['completion_pct']}%"
            f" ({m['completed_tasks']}/{m['total_tasks']}) |"
        )
        md.append(
            f"| **Success Criteria Met** |"
            f" {m['criteria_met']}/{m['criteria_total']}"
            f" ({m['criteria_pct']}%) |"
        )
        md.append(f"| **Success Score** | {score_text} |")
        md.append(f"| **Days Remaining** | {days_text} |")
        md.append(f"| **Task Groups** | {m['total_groups']} |")
        md.append(
            f"| **Total Comments** |"
            f" {m['total_comments']}"
            f" (Internal: {m['internal_comments']}) |"
        )
        md.append(f"| **Participants** | {m['total_participants']} |")
        md.append("")

        # Charts
        if charts.get("progress_gauge"):
            md.append("### Overall Progress\n")
            md.append(
                f'<div align="center">'
                f'<img src="{charts["progress_gauge"]}"'
                f' alt="Progress Gauge" width="600"/></div>\n'
            )

        if charts.get("task_status"):
            md.append("### Task Status Distribution\n")
            md.append(
                f'<div align="center">'
                f'<img src="{charts["task_status"]}"'
                f' alt="Task Status" width="500"/></div>\n'
            )

        if charts.get("success_criteria"):
            md.append("### Success Criteria Achievement\n")
            md.append(
                f'<div align="center">'
                f'<img src="{charts["success_criteria"]}"'
                f' alt="Success Criteria" width="600"/></div>\n'
            )

        if charts.get("timeline"):
            md.append("### POC Timeline\n")
            md.append(
                f'<div align="center">'
                f'<img src="{charts["timeline"]}"'
                f' alt="Timeline" width="700"/></div>\n'
            )

        if charts.get("workload"):
            md.append("### Team Workload\n")
            md.append(
                f'<div align="center">'
                f'<img src="{charts["workload"]}"'
                f' alt="Workload" width="600"/></div>\n'
            )

        if charts.get("activity"):
            md.append("### Activity Over Time\n")
            md.append(
                f'<div align="center">'
                f'<img src="{charts["activity"]}"'
                f' alt="Activity" width="600"/></div>\n'
            )

        md.append("---\n")

        # =============================================
        #  POC OVERVIEW
        # =============================================
        if (
            self.poc.executive_summary
            or self.poc.description
            or self.poc.objectives
        ):
            md.append(
                f'<h2 style="color: {self._primary_hex};">'
                f"POC Overview</h2>\n"
            )
            if self.poc.executive_summary:
                md.append("### Executive Summary\n")
                md.append(f"{self.poc.executive_summary}\n")
            if self.poc.description:
                md.append("### Description\n")
                md.append(f"{self.poc.description}\n")
            if self.poc.objectives:
                md.append("### Objectives\n")
                md.append(f"{self.poc.objectives}\n")

        # Products
        if self.poc.products:
            md.append("### Products\n")
            for product in self.poc.products:
                md.append(f"- **{product.name}**")
            md.append("")

        # =============================================
        #  SUCCESS CRITERIA TABLE
        # =============================================
        criteria = self.poc.success_criteria
        if criteria:
            md.append(
                f'<h2 style="color: {self._primary_hex};">'
                f"Success Criteria</h2>\n"
            )
            md.append(
                "| # | Criteria | Target | Achieved" " | Importance | Met |"
            )
            md.append(
                "|---|----------|--------|----------" "|------------|-----|"
            )
            for i, c in enumerate(criteria, 1):
                stars = "★" * (c.importance_level or 3)
                met = "✅" if c.is_met else "❌"
                md.append(
                    f"| {i} | {c.title}"
                    f" | {c.target_value or 'N/A'}"
                    f" | {c.achieved_value or 'N/A'}"
                    f" | {stars} | {met} |"
                )
            md.append("")

        # =============================================
        #  TASKS
        # =============================================
        tasks = self._get_all_tasks()
        if tasks:
            md.append(
                f'<h2 style="color: {self._primary_hex};">' f"Tasks</h2>\n"
            )
            for task in tasks:
                md.extend(self._md_render_task(task))

            md.append("---\n")

        # =============================================
        #  TASK GROUPS
        # =============================================
        groups = self._get_all_task_groups()
        if groups:
            md.append(
                f'<h2 style="color: {self._primary_hex};">'
                f"Task Groups</h2>\n"
            )
            for group in groups:
                status_val = (
                    group.status.value if group.status else "not_started"
                )
                status_emoji = {
                    "not_started": "⚪",
                    "in_progress": "🔵",
                    "completed": "✅",
                    "blocked": "🔴",
                }
                emoji = status_emoji.get(status_val, "⚪")
                md.append(
                    f"### 📁 {group.title} — {emoji}"
                    f" {status_val.replace('_', ' ').title()}\n"
                )

                if group.description:
                    md.append(f"{group.description}\n")

                # Group resources
                group_resources = (
                    self.db.query(Resource)
                    .filter(Resource.poc_task_group_id == group.id)
                    .all()
                )
                if group_resources:
                    md.append("**Group Resources:**")
                    for res in group_resources:
                        md.append(
                            f"- **{res.title}**"
                            f" ({res.resource_type.value})"
                        )
                        if res.description:
                            md.append(f"  - {res.description}")
                        if res.resource_type.value == "LINK" and res.content:
                            md.append(f"  - Link: {res.content}")
                    md.append("")

                # Group comments
                group_comments = (
                    self.db.query(Comment)
                    .filter(Comment.poc_task_group_id == group.id)
                    .order_by(Comment.created_at.desc())
                    .limit(5)
                    .all()
                )
                if group_comments:
                    md.append("**Group Comments:**")
                    for c in reversed(group_comments):
                        author = self._comment_author_name(c)
                        dt = c.created_at.strftime("%Y-%m-%d %H:%M")
                        vis = "🔒 Internal" if c.is_internal else "👁️ Public"
                        md.append(f"- **{author}** ({vis}) — {dt}")
                        md.append(f"  {c.content}")
                    md.append("")

                # Tasks in group
                group_tasks = self._get_group_tasks(group)
                if group_tasks:
                    md.append("**Tasks in this group:**\n")
                    for task in group_tasks:
                        md.extend(self._md_render_task(task, level=4))

                md.append("---\n")

        # =============================================
        #  POC-LEVEL RESOURCES
        # =============================================
        poc_resources = self._get_poc_resources()
        if poc_resources:
            md.append(
                f'<h2 style="color: {self._primary_hex};">'
                f"POC Resources</h2>\n"
            )
            for resource in poc_resources:
                md.append(f"### {resource.title}\n")
                md.append(f"**Type:** {resource.resource_type.value}\n")
                if resource.description:
                    md.append(f"{resource.description}\n")
                if resource.resource_type.value == "LINK" and resource.content:
                    md.append(f"**Link:** {resource.content}\n")
                elif (
                    resource.resource_type.value == "CODE" and resource.content
                ):
                    md.append(f"```\n{resource.content}\n```\n")
                elif resource.content:
                    md.append(f"{resource.content}\n")
                md.append("")

        # =============================================
        #  PARTICIPANTS
        # =============================================
        participants = self._get_participants()
        if participants:
            md.append(
                f'<h2 style="color: {self._primary_hex};">'
                f"Participants</h2>\n"
            )
            md.append("| Name | Email | Role | Joined |")
            md.append("|------|-------|------|--------|")
            for p in participants:
                if p.user:
                    name = p.user.full_name or p.user.email
                    email = p.user.email
                    role_parts = []
                    if p.is_sales_engineer:
                        role_parts.append("Sales Engineer")
                    if p.is_customer:
                        role_parts.append("Customer")
                    role = (
                        ", ".join(role_parts) if role_parts else "Participant"
                    )
                    joined = (
                        p.joined_at.strftime("%Y-%m-%d")
                        if p.joined_at
                        else "N/A"
                    )
                    md.append(f"| {name} | {email}" f" | {role} | {joined} |")
            md.append("")

        # =============================================
        #  POC-LEVEL COMMENTS
        # =============================================
        poc_comments = (
            self.db.query(Comment)
            .filter(
                Comment.poc_id == self.poc.id,
                Comment.poc_task_id.is_(None),
                Comment.poc_task_group_id.is_(None),
            )
            .order_by(Comment.created_at.desc())
            .limit(10)
            .all()
        )
        if poc_comments:
            md.append(
                f'<h2 style="color: {self._primary_hex};">'
                f"Recent POC Comments</h2>\n"
            )
            for c in reversed(poc_comments):
                author = self._comment_author_name(c)
                dt = c.created_at.strftime("%Y-%m-%d %H:%M")
                vis = "🔒 Internal" if c.is_internal else "👁️ Public"
                md.append(f"- **{author}** ({vis}) — {dt}")
                md.append(f"  {c.content}")
            md.append("")

        # Write output
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md))

        return output_path

    def _md_render_task(self, task, level: int = 3) -> List[str]:
        """Render a single task to Markdown lines."""
        lines: List[str] = []
        status_emoji = {
            "not_started": "⚪",
            "in_progress": "🔵",
            "completed": "✅",
            "blocked": "🔴",
        }
        emoji = status_emoji.get(task.status.value, "⚪")
        heading = "#" * level

        lines.append(f"{heading} {task.title}\n")
        lines.append(
            f"**Status:** {emoji}"
            f" {task.status.value.replace('_', ' ').title()}\n"
        )

        if task.description:
            lines.append(f"{task.description}\n")

        # Dates
        date_parts = []
        if task.start_date:
            date_parts.append(f"**Start:** {task.start_date}")
        if task.due_date:
            date_parts.append(f"**Due:** {task.due_date}")
        if task.completed_at:
            date_parts.append(f"**Completed:** {task.completed_at}")
        if date_parts:
            lines.append(" | ".join(date_parts) + "\n")

        # Assignees
        assignees = (
            self.db.query(POCTaskAssignee)
            .filter(POCTaskAssignee.poc_task_id == task.id)
            .all()
        )
        if assignees:
            lines.append("**Assigned to:**")
            for a in assignees:
                if a.participant and a.participant.user:
                    name = (
                        a.participant.user.full_name
                        or a.participant.user.email
                    )
                    lines.append(
                        f"- 👤 {name}" f" ({a.participant.user.email})"
                    )
            lines.append("")

        # Resources
        task_resources = (
            self.db.query(Resource)
            .filter(Resource.poc_task_id == task.id)
            .all()
        )
        if task_resources:
            lines.append("**Resources:**")
            for res in task_resources:
                lines.append(
                    f"- **{res.title}**" f" ({res.resource_type.value})"
                )
                if res.description:
                    lines.append(f"  - {res.description}")
                if res.resource_type.value == "LINK" and res.content:
                    lines.append(f"  - Link: {res.content}")
            lines.append("")

        # Success criteria linked to this task
        if hasattr(task, "task_criteria") and task.task_criteria:
            lines.append("**Success Criteria:**")
            for tc in task.task_criteria:
                sc = (
                    tc.success_criteria
                    if hasattr(tc, "success_criteria")
                    else None
                )
                if sc:
                    met = "✅" if sc.is_met else "❌"
                    lines.append(f"- {met} {sc.title}")
            lines.append("")

        # Comments
        comments = (
            self.db.query(Comment)
            .filter(Comment.poc_task_id == task.id)
            .order_by(Comment.created_at.desc())
            .limit(5)
            .all()
        )
        if comments:
            lines.append("**Latest Comments:**")
            for c in reversed(comments):
                author = self._comment_author_name(c)
                dt = c.created_at.strftime("%Y-%m-%d %H:%M")
                vis = "🔒 Internal" if c.is_internal else "👁️ Public"
                lines.append(f"- **{author}** ({vis}) — {dt}")
                lines.append(f"  {c.content}")
            lines.append("")

        lines.append("")
        return lines

    # ------------------------------------------------------------------ #
    #  Word Generation                                                     #
    # ------------------------------------------------------------------ #

    def generate_word(self, output_path: str):
        """Generate Word document (basic format)."""
        doc = Document()

        # Customer logo
        if self.poc.customer_logo_url:
            logo_path = self._get_logo_path(self.poc.customer_logo_url)
            if logo_path and os.path.exists(logo_path):
                try:
                    logo_paragraph = doc.add_paragraph()
                    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = logo_paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(2))
                except Exception:
                    pass

        # Customer name
        customer_para = doc.add_paragraph()
        customer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        customer_run = customer_para.add_run(self.poc.customer_company_name)
        customer_run.font.size = Pt(18)
        customer_run.font.bold = True
        customer_run.font.color.rgb = RGBColor(51, 51, 51)

        doc.add_paragraph()

        # Title
        title = doc.add_heading(self.poc.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # POC Details
        doc.add_heading("POC Details", level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"

        details = [
            ("Customer", self.poc.customer_company_name),
            ("Status", self.poc.status.value),
            (
                "Start Date",
                (str(self.poc.start_date) if self.poc.start_date else "N/A"),
            ),
            (
                "End Date",
                (str(self.poc.end_date) if self.poc.end_date else "N/A"),
            ),
            (
                "Success Score",
                (
                    f"{self.poc.overall_success_score}/100"
                    if self.poc.overall_success_score
                    else "N/A"
                ),
            ),
        ]
        for i, (label, value) in enumerate(details):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        # Description
        if self.poc.description:
            doc.add_heading("Description", level=1)
            doc.add_paragraph(self.poc.description)

        # Success Criteria
        if self.poc.success_criteria:
            doc.add_heading("Success Criteria", level=1)
            criteria_table = doc.add_table(
                rows=len(self.poc.success_criteria) + 1, cols=4
            )
            criteria_table.style = "Light Grid Accent 1"
            headers = ["Criteria", "Target", "Achieved", "Met"]
            for i, header in enumerate(headers):
                criteria_table.rows[0].cells[i].text = header
            for i, criteria in enumerate(self.poc.success_criteria, 1):
                criteria_table.rows[i].cells[0].text = criteria.title
                criteria_table.rows[i].cells[1].text = (
                    criteria.target_value or "N/A"
                )
                criteria_table.rows[i].cells[2].text = (
                    criteria.achieved_value or "N/A"
                )
                criteria_table.rows[i].cells[3].text = (
                    "Yes" if criteria.is_met else "No"
                )

        # Tasks
        if self.poc.poc_tasks:
            doc.add_heading("Tasks", level=1)
            for task in self.poc.poc_tasks:
                doc.add_heading(task.title, level=2)
                if task.description:
                    doc.add_paragraph(task.description)
                doc.add_paragraph(f"Status: {task.status.value}")
                if task.success_level:
                    doc.add_paragraph(
                        f"Success Level: {task.success_level}/100"
                    )

        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------------ #
    #  Utilities                                                           #
    # ------------------------------------------------------------------ #

    def _get_logo_path(self, logo_url: str) -> Optional[str]:
        """Resolve logo URL to a local file path."""
        from app.config import Settings

        settings = Settings()
        relative_path = logo_url.lstrip("/")

        if relative_path.startswith("uploads/"):
            relative_path = relative_path[8:]

        logo_path = Path(settings.UPLOAD_DIR) / relative_path
        return str(logo_path) if logo_path.exists() else None
